"""Phase 107's gate: does a report only restate what a run produced?

Three things are worth testing here and they are not the obvious ones.

**The stamp is read back out of the file, not asserted from the code that wrote
it.** T107.3 asks that every report carry the screening disclaimer and the model
version. A test that checks the constant is present in the module proves nothing
about the document a reader opens. `python-docx` parses the written `.docx`
back, and the assertions run against that.

**The experiment report may not round differently from the tables engine.** Every
metric cell is compared twice: once against `tables.format_value`, which is the
single place rounding is allowed to happen, and once against an independently
written `format(value, ".3f")`. Two agreeing implementations are a real check;
comparing the report against the function it called would not be.

**The contribution decomposition has to add up.** For a linear model the report
claims the terms plus the intercept are the log-odds. That claim is checked
against `decision_function` rather than taken on trust -- a ranking of numbers
that does not sum to the decision is a plausible-looking figure with nothing
behind it, which is the failure mode this project exists to avoid.

Everything needing `dataset/` or the saved bundle skips: both are gitignored, so
CI runs the experiment-report half and the structural half of this file. The
per-recording half is run locally, where the corpus is.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import numpy as np
import pytest

from src.inference.predictor import DISCLAIMER, ModelBundle
from src.reporting.sample_report import (
    CONTRIBUTION_UNAVAILABLE,
    REPORTED_METRICS,
    ReportError,
    batch_export,
    feature_contributions,
    render_experiment_report,
    report_for_recording,
)

PROJECT_ROOT = Path(__file__).resolve().parents[1]
EXP_DIR = PROJECT_ROOT / "outputs" / "06_binary_results" / "EXP-A1"
BUNDLE_DIR = PROJECT_ROOT / "models_saved" / "binary" / "final"
PHYSIONET = PROJECT_ROOT / "dataset" / "archive (3)" / "training-a"

#: How far the re-summed contribution terms may sit from `decision_function`.
#:
#: Adding 138 products in a different order from the BLAS call inside
#: `decision_function` changes the result in its last bits -- the same
#: accumulation-order effect Phase 106 measured between a single record and a
#: batch, and the same one the Part VII re-run found in `brier` and `ece`. This
#: is a re-summation of the identical terms, so the tolerance is relative and
#: tight; it is not a licence for the decomposition to be approximately right.
LOGIT_RTOL = 1e-9


# ---------------------------------------------------------------------------
# the disclaimer and the language, which have no dataset dependency
# ---------------------------------------------------------------------------


def test_the_report_module_reuses_the_predictor_disclaimer_rather_than_copying_it() -> None:
    from src.reporting import sample_report

    assert sample_report.DISCLAIMER is DISCLAIMER
    text = (PROJECT_ROOT / "src" / "reporting" / "sample_report.py").read_text(encoding="utf-8")
    assert "Academic screening and decision-support prototype" not in text, (
        "the disclaimer text is written out a second time here; it must be imported "
        "so a change in one place cannot leave two versions in circulation"
    )


def test_the_unavailable_reason_refuses_to_substitute_global_importance() -> None:
    lowered = CONTRIBUTION_UNAVAILABLE.lower()
    assert "shap" in lowered
    assert "global feature importance is deliberately not substituted" in lowered


def test_the_reported_metric_set_is_the_one_research_rule_six_requires() -> None:
    names = [name for name, _label in REPORTED_METRICS]
    for required in ("sensitivity", "specificity", "balanced_accuracy", "f1", "roc_auc"):
        assert required in names, required + " is not reported"
    assert names.index("sensitivity") < names.index("accuracy"), (
        "accuracy leads the table; rule 6 makes sensitivity and balanced accuracy the headline"
    )


# ---------------------------------------------------------------------------
# T107.2 -- contributions, and the estimators that cannot have them
# ---------------------------------------------------------------------------


def _bundle_for(pipeline: Any, feature_names: tuple[str, ...]) -> ModelBundle:
    return ModelBundle(
        task="binary",
        pipeline=pipeline,
        feature_names=feature_names,
        manifest={"model_id": "test"},
        path=Path("."),
    )


@pytest.fixture(scope="module")
def synthetic_frame() -> tuple[np.ndarray, np.ndarray, tuple[str, ...]]:
    rng = np.random.default_rng(42)
    x = rng.normal(size=(80, 6))
    y = (x[:, 0] + 0.5 * x[:, 1] > 0).astype(int)
    return x, y, tuple("f" + str(index) for index in range(6))


def test_a_forest_gets_the_unavailable_reason_and_no_numbers(
    synthetic_frame: tuple[np.ndarray, np.ndarray, tuple[str, ...]],
) -> None:
    from sklearn.ensemble import RandomForestClassifier
    from sklearn.pipeline import Pipeline
    from sklearn.preprocessing import StandardScaler

    x, y, names = synthetic_frame
    pipeline = Pipeline(
        [
            ("scale", StandardScaler()),
            ("clf", RandomForestClassifier(n_estimators=5, random_state=42)),
        ]
    )
    pipeline.fit(x, y)

    contributions, reason = feature_contributions(_bundle_for(pipeline, names), x[0])
    assert contributions == []
    assert reason == CONTRIBUTION_UNAVAILABLE


def test_a_multiclass_linear_model_gets_the_unavailable_reason() -> None:
    from sklearn.linear_model import LogisticRegression
    from sklearn.pipeline import Pipeline
    from sklearn.preprocessing import StandardScaler

    rng = np.random.default_rng(42)
    x = rng.normal(size=(90, 4))
    y = rng.integers(0, 3, size=90)
    pipeline = Pipeline([("scale", StandardScaler()), ("clf", LogisticRegression(max_iter=500))])
    pipeline.fit(x, y)

    contributions, reason = feature_contributions(_bundle_for(pipeline, ("a", "b", "c", "d")), x[0])
    assert contributions == []
    assert reason == CONTRIBUTION_UNAVAILABLE, (
        "one contribution is not defined without naming a class, and inventing one "
        "would be a number with no meaning"
    )


def test_the_contribution_terms_and_the_intercept_are_the_log_odds(
    synthetic_frame: tuple[np.ndarray, np.ndarray, tuple[str, ...]],
) -> None:
    """The claim the report prints, checked against the model's own decision."""
    from sklearn.linear_model import LogisticRegression
    from sklearn.pipeline import Pipeline
    from sklearn.preprocessing import StandardScaler

    x, y, names = synthetic_frame
    pipeline = Pipeline([("scale", StandardScaler()), ("clf", LogisticRegression(max_iter=500))])
    pipeline.fit(x, y)
    bundle = _bundle_for(pipeline, names)

    # `top` is the full width, so the terms are the whole decomposition.
    contributions, reason = feature_contributions(bundle, x[3], top=len(names))
    assert reason is None
    assert len(contributions) == len(names)

    intercept = float(pipeline.steps[-1][1].intercept_[0])
    summed = sum(item.contribution for item in contributions) + intercept
    expected = float(pipeline.decision_function(x[3].reshape(1, -1))[0])
    assert np.isclose(summed, expected, rtol=LOGIT_RTOL, atol=0.0), (
        "the terms do not sum to the decision: " + repr(summed) + " vs " + repr(expected)
    )


def test_contributions_are_ordered_by_absolute_size(
    synthetic_frame: tuple[np.ndarray, np.ndarray, tuple[str, ...]],
) -> None:
    from sklearn.linear_model import LogisticRegression
    from sklearn.pipeline import Pipeline
    from sklearn.preprocessing import StandardScaler

    x, y, names = synthetic_frame
    pipeline = Pipeline([("scale", StandardScaler()), ("clf", LogisticRegression(max_iter=500))])
    pipeline.fit(x, y)
    contributions, _reason = feature_contributions(_bundle_for(pipeline, names), x[5], top=4)

    magnitudes = [abs(item.contribution) for item in contributions]
    assert magnitudes == sorted(magnitudes, reverse=True)
    assert len(contributions) == 4


# ---------------------------------------------------------------------------
# T107.4 -- the experiment report, which needs only committed CSVs
# ---------------------------------------------------------------------------


@pytest.fixture(scope="module")
def experiment_report(tmp_path_factory: pytest.TempPathFactory) -> Path:
    if not (EXP_DIR / "aggregate_metrics.csv").is_file():
        pytest.skip("EXP-A1 aggregate_metrics.csv is not in this checkout")
    target = tmp_path_factory.mktemp("experiment") / "EXP-A1.docx"
    return render_experiment_report(EXP_DIR, target)


def _table_headed(path: Path, first_cell: str) -> list[list[str]]:
    from docx import Document

    for grid in Document(str(path)).tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in grid.rows]
        if rows and rows[0] and rows[0][0] == first_cell:
            return rows
    raise AssertionError("no table whose first header cell is " + repr(first_cell))


def test_every_rendered_metric_equals_its_source_cell(experiment_report: Path) -> None:
    """Read the document back and compare each cell against the CSV.

    Checked twice on purpose: against `tables.format_value`, the project's single
    rounding authority, and against an independently written three-decimal
    format. Comparing only against the function the writer called would test
    nothing.
    """
    import pandas as pd

    from src.reporting.tables import format_value

    frame = pd.read_csv(EXP_DIR / "aggregate_metrics.csv")
    rows = _table_headed(experiment_report, "Metric")
    models = rows[0][1:]
    assert models == [str(value) for value in frame["model_id"].tolist()]

    labels = {label: name for name, label in REPORTED_METRICS}
    checked = 0
    for row in rows[1:]:
        column = labels[row[0]]
        for index, cell in enumerate(row[1:]):
            mean = frame[column + "_mean"].iloc[index]
            sd = frame[column + "_sd"].iloc[index]
            assert cell == format_value(mean, "metric") + " +/- " + format_value(sd, "metric")
            assert cell == format(float(mean), ".3f") + " +/- " + format(float(sd), ".3f")
            checked += 1
    assert checked == (len(rows) - 1) * len(models)
    assert checked >= 30, "too few cells checked to mean anything: " + str(checked)


def test_the_experiment_report_carries_the_disclaimer_and_the_run_provenance(
    experiment_report: Path,
) -> None:
    import json as _json

    from docx import Document

    document = Document(str(experiment_report))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    assert paragraphs.count(DISCLAIMER) >= 2, "the disclaimer must open and close the report"
    assert paragraphs[1] == DISCLAIMER, "the disclaimer must precede every number"

    manifest = _json.loads((EXP_DIR / "run_manifest.json").read_text(encoding="utf-8"))
    run_info = dict(_table_headed(experiment_report, "Experiment directory"))  # type: ignore[arg-type]
    assert run_info["Run id"] == str(manifest["run_id"])
    assert run_info["Seed"] == str(manifest["seed"])
    assert run_info["Git commit"] == str(manifest["git"]["commit"])


def test_the_experiment_report_records_the_digest_of_every_file_it_read(
    experiment_report: Path,
) -> None:
    from src.reporting.tables import source_fingerprint

    sources = dict(_table_headed(experiment_report, "aggregate_metrics.csv sha256"))  # type: ignore[arg-type]
    for name in ("aggregate_metrics.csv", "per_fold_metrics.csv", "run_manifest.json"):
        key = name + " sha256"
        assert key in sources, key + " is not recorded"
        expected = str(source_fingerprint(EXP_DIR / name)["sha256"])[:16]
        assert sources[key] == expected


def test_a_directory_without_aggregate_metrics_is_refused(tmp_path: Path) -> None:
    with pytest.raises(ReportError):
        render_experiment_report(tmp_path, tmp_path / "out.docx")


# ---------------------------------------------------------------------------
# T107.5 -- batch export
# ---------------------------------------------------------------------------


def test_an_empty_batch_is_refused_rather_than_written(tmp_path: Path) -> None:
    with pytest.raises(ReportError):
        batch_export([], tmp_path)


# ---------------------------------------------------------------------------
# T107.1 / T107.2 / T107.3 -- the per-recording report, against a real recording
# ---------------------------------------------------------------------------


@pytest.fixture(scope="module")
def rendered_sample(tmp_path_factory: pytest.TempPathFactory) -> tuple[Path, Any]:
    if not (BUNDLE_DIR / "model.joblib").is_file():
        pytest.skip("no saved binary model in this checkout (gitignored)")
    if not PHYSIONET.is_dir():
        pytest.skip("dataset/ is not present in this checkout")
    recordings = sorted(PHYSIONET.glob("*.wav"))
    if not recordings:
        pytest.skip("no PhysioNet training-a recordings on disk")

    out = tmp_path_factory.mktemp("sample")
    return report_for_recording(recordings[0], out / "sample.docx", figures_dir=out / "figures")


def test_the_report_stamps_the_disclaimer_before_and_after_every_number(
    rendered_sample: tuple[Path, Any],
) -> None:
    from docx import Document

    path, _result = rendered_sample
    paragraphs = [p.text.strip() for p in Document(str(path)).paragraphs if p.text.strip()]
    assert paragraphs[1] == DISCLAIMER, (
        "the disclaimer is not the first thing after the title; a disclaimer that "
        "follows the result is a disclaimer read after the decision"
    )
    assert paragraphs[-1] == DISCLAIMER


def test_the_report_stamps_the_model_version_from_the_bundle_manifest(
    rendered_sample: tuple[Path, Any],
) -> None:
    path, result = rendered_sample
    block = dict(_table_headed(path, "Model id"))  # type: ignore[arg-type]
    assert block["Model id"] == str(result.model["model_id"])
    assert block["Estimator"] == str(result.model["estimator_class"])
    assert block["Saved at"] == str(result.model["saved_at"])
    assert block["Records fitted"] == str(result.model["n_records_fitted"])
    assert block["Features"] == "138"


def test_the_report_prints_the_prediction_the_predictor_returned(
    rendered_sample: tuple[Path, Any],
) -> None:
    path, result = rendered_sample
    screening = dict(_table_headed(path, "Predicted class"))  # type: ignore[arg-type]
    assert screening["Predicted class"] == result.predicted_class
    assert screening["Confidence"] == format(result.confidence, ".4f")

    probabilities = _table_headed(path, "Class")
    rendered = {row[0]: row[1] for row in probabilities[1:]}
    assert rendered == {name: format(value, ".4f") for name, value in result.probabilities.items()}


def test_the_report_states_the_operating_point_is_a_plain_argmax(
    rendered_sample: tuple[Path, Any],
) -> None:
    from docx import Document

    path, _result = rendered_sample
    text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    assert "no in-fold selected threshold" in text, (
        "a reader who knows T50.4 exists will assume a tuned operating point was "
        "applied unless the report says otherwise"
    )


def test_the_report_embeds_the_waveform_and_the_spectrogram(
    rendered_sample: tuple[Path, Any],
) -> None:
    from docx import Document

    path, _result = rendered_sample
    assert len(Document(str(path)).inline_shapes) == 2


def test_the_report_uses_screening_language_only(rendered_sample: tuple[Path, Any]) -> None:
    from docx import Document

    path, _result = rendered_sample
    document = Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs)
    for grid in document.tables:
        for row in grid.rows:
            text += "\n" + "\n".join(cell.text for cell in row.cells)

    lowered = text.lower()
    assert "not a diagnostic device" in lowered
    for forbidden in ("diagnosis of", "the patient", "case study", "treatment plan"):
        assert forbidden not in lowered, "clinical language in a screening report: " + forbidden


def test_the_report_renders_the_vector_that_was_actually_scored() -> None:
    """No second computation between the prediction and the picture of it.

    The report draws the signal and decomposes the vector handed back by
    `predict_recording`, so re-scoring that vector through the same bundle
    reproduces the probability **exactly** -- the same call on the same row is
    the same BLAS call, and there is no batch/single split to explain a
    difference away.
    """
    if not (BUNDLE_DIR / "model.joblib").is_file():
        pytest.skip("no saved binary model in this checkout (gitignored)")
    if not PHYSIONET.is_dir():
        pytest.skip("dataset/ is not present in this checkout")
    recordings = sorted(PHYSIONET.glob("*.wav"))
    if not recordings:
        pytest.skip("no PhysioNet training-a recordings on disk")

    from threadpoolctl import threadpool_limits

    from src.inference.predictor import predict_recording

    result, detail = predict_recording(recordings[0], with_detail=True)
    assert detail.vector.shape == (len(detail.feature_names),)
    assert detail.feature_names == detail.bundle.feature_names

    with threadpool_limits(1):
        again = detail.bundle.pipeline.predict_proba(detail.vector.reshape(1, -1))[0]
    assert float(again[result.predicted_index]) == result.confidence


def test_batch_export_round_trips_and_is_strict_json(
    rendered_sample: tuple[Path, Any], tmp_path: Path
) -> None:
    import pandas as pd

    _path, result = rendered_sample
    written = batch_export([result], tmp_path)

    frame = pd.read_csv(written["csv"], comment="#")
    assert len(frame) == 1
    assert frame["predicted_class"].iloc[0] == result.predicted_class
    assert written["csv"].read_text(encoding="utf-8").splitlines()[0] == "# " + DISCLAIMER

    raw = written["json"].read_text(encoding="utf-8")
    for token in ("NaN", "Infinity"):
        assert token not in raw, token + " is not JSON and no strict parser will read it"
    payload = json.loads(raw)
    assert payload["disclaimer"] == DISCLAIMER
    assert payload["n_records"] == 1
    assert payload["predictions"][0]["predicted_class"] == result.predicted_class
