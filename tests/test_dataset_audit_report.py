"""DA-01 .. DA-09 completeness gate (T21.7).

Two things are checked. First, that all nine audit artifacts exist and are
non-empty -- a generator that writes an empty CSV has not succeeded. Second,
that the narrative report actually *says* the three things T21.2 -- T21.4
require it to say, rather than merely existing.

The .docx is read back and its text searched. A report that was generated but
whose content silently lost a section would pass a mere existence check.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pytest

from src.data_loader import inventory as inv

# ---------------------------------------------------------------------------
# fixtures
# ---------------------------------------------------------------------------


@pytest.fixture(scope="module")
def audit_dir() -> Path:
    from src.utils.config import load_config

    return Path(load_config("paths").require("outputs.dataset_audit"))


@pytest.fixture(scope="module")
def master() -> Any:
    from src.data_loader.master import load_master

    return load_master()


@pytest.fixture(scope="module")
def report_text(audit_dir: Path) -> str:
    from docx import Document

    path = audit_dir / "dataset_audit_report.docx"
    if not path.is_file():
        pytest.fail("DA-09 has not been generated -- run scripts/01_run_dataset_audit.py")
    document = Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ===========================================================================
# T21.7 -- all nine artifacts exist and are non-empty
# ===========================================================================


def test_nine_artifacts_are_registered() -> None:
    ids = [row[0] for row in inv.AUDIT_ARTIFACTS]
    assert ids == ["DA-0" + str(n) for n in range(1, 10)]


@pytest.mark.needs_data
def test_every_da_artifact_exists_and_is_non_empty(audit_dir: Path) -> None:
    for evidence_id, filename, _ in inv.AUDIT_ARTIFACTS:
        path = audit_dir / filename
        assert path.is_file(), evidence_id + " missing: " + str(path)
        assert path.stat().st_size > 0, evidence_id + " is empty"


@pytest.mark.needs_data
def test_every_da_csv_has_rows_not_just_a_header(audit_dir: Path) -> None:
    """An all-header CSV is a generator that ran and produced nothing."""
    import pandas as pd

    for evidence_id, filename, _ in inv.AUDIT_ARTIFACTS:
        if not filename.endswith(".csv"):
            continue
        frame = pd.read_csv(audit_dir / filename, keep_default_na=False)
        assert len(frame) > 0, evidence_id + " (" + filename + ") has no rows"


@pytest.mark.needs_data
def test_all_nine_appear_in_the_evidence_index() -> None:
    """T21.6 -- registered, and registered as present rather than missing."""
    from src.utils.evidence import read_evidence

    rows = {row["evidence_id"]: row for row in read_evidence()}
    for evidence_id, _, _ in inv.AUDIT_ARTIFACTS:
        assert evidence_id in rows, evidence_id + " is not in the evidence index"
        assert rows[evidence_id]["status"] == "ok", evidence_id


# ===========================================================================
# T21.1 -- DA-01 contents
# ===========================================================================


@pytest.mark.needs_data
def test_inventory_covers_four_datasets_with_the_audited_counts(master: Any) -> None:
    frame = inv.build_inventory(master)
    assert list(frame.columns) == list(inv.INVENTORY_COLUMNS)
    assert list(frame["dataset_source"]) == ["D1", "D2", "D3", "D4"]

    expected = {
        # dataset: (total_files, usable_files, native_fs)
        "D1": (3541, 3240, 2000),
        "D2": (176, 124, 44100),
        "D3": (656, 461, 4000),
        "D4": (3163, 3163, 4000),
    }
    for row in frame.itertuples(index=False):
        total, usable, fs = expected[str(row.dataset_source)]
        assert int(row.total_files) == total, row.dataset_source
        assert int(row.usable_files) == usable, row.dataset_source
        assert int(row.original_fs) == fs, row.dataset_source
        assert int(row.target_fs) == 2000
        assert float(row.total_hours) > 0
        assert str(row.role).strip()
        assert str(row.folder).startswith("dataset/")

    assert int(frame["total_files"].sum()) == 7536
    assert int(frame["usable_files"].sum()) == 6988


@pytest.mark.needs_data
def test_inventory_never_merges_two_label_spaces_into_one_class_list(
    master: Any,
) -> None:
    """Rule 4, at the inventory level.

    D4 owns two tasks, so its class list names both; D2 and D3 own one each and
    must never appear in the same cell.
    """
    frame = inv.build_inventory(master).set_index("dataset_source")
    assert frame.loc["D2", "classes"].startswith("pascal_a:")
    assert "pascal_b" not in frame.loc["D2", "classes"]
    assert frame.loc["D3", "classes"].startswith("pascal_b:")
    assert "pascal_a" not in frame.loc["D3", "classes"]
    assert "circor_murmur:" in frame.loc["D4", "classes"]
    assert "circor_outcome:" in frame.loc["D4", "classes"]


# ===========================================================================
# T21.2 / T21.3 / T21.4 -- the report says what it must say
# ===========================================================================


def test_both_count_discrepancies_are_recorded() -> None:
    subjects = {d.subject for d in inv.DISCREPANCIES}
    assert any("CirCor" in s for s in subjects)
    assert any("PhysioNet" in s for s in subjects)
    assert {d.task for d in inv.DISCREPANCIES} == {"T21.2", "T21.3"}


def test_the_pascal_a_subject_limitation_is_recorded() -> None:
    texts = [item.subject + " " + item.explanation for item in inv.LIMITATIONS]
    assert any("no recoverable subject IDs" in t for t in texts)
    assert any("RECORD-LEVEL ONLY" in t for t in texts)


@pytest.mark.needs_data
def test_report_states_the_circor_count_discrepancy(report_text: str) -> None:
    assert "1,568" in report_text
    assert "5,272" in report_text
    assert "942" in report_text and "3,163" in report_text
    assert "hidden validation and test sets" in report_text


@pytest.mark.needs_data
def test_report_states_the_physionet_count_discrepancy(report_text: str) -> None:
    assert "3,126" in report_text
    assert "3,240" in report_text
    assert "301" in report_text
    assert "byte-identical" in report_text


@pytest.mark.needs_data
def test_report_states_the_pascal_a_subject_limitation(report_text: str) -> None:
    assert "no recoverable subject IDs" in report_text
    assert "RECORD-LEVEL ONLY" in report_text


@pytest.mark.needs_data
def test_report_carries_the_screening_disclaimer(report_text: str) -> None:
    """Rule 7 -- screening language, in every generated report."""
    assert "screening and decision-support" in report_text
    assert "not a diagnostic device" in report_text
    for forbidden in ("diagnoses the patient", "replaces a doctor", "prescribes"):
        assert forbidden not in report_text


@pytest.mark.needs_data
def test_report_states_that_label_spaces_are_never_merged(report_text: str) -> None:
    assert "never merged" in report_text
    assert "five separate tasks" in report_text


@pytest.mark.needs_data
def test_report_numbers_agree_with_the_csvs_they_describe(
    master: Any, audit_dir: Path
) -> None:
    """No hand-typed figure survives a regeneration -- so check a few.

    The point of generating the .docx from code is that its prose cannot drift
    from the CSVs. This asserts that it has not.
    """
    import pandas as pd

    frame = inv.build_inventory(master)
    sections = inv.build_report_sections(master, frame, audit_dir=audit_dir)
    text = "\n".join(p for _, paragraphs in sections for p in paragraphs)

    duplicates = pd.read_csv(audit_dir / "duplicate_report.csv", keep_default_na=False)
    assert str(len(duplicates)) + " rows" in text

    class_dist = pd.read_csv(audit_dir / "class_distribution.csv")
    supervised = class_dist[class_dist["scope"] == "supervised"]
    row = supervised[
        (supervised["task"] == "binary") & (supervised["class"] == "abnormal")
    ].iloc[0]
    assert "abnormal " + str(int(row["n_records"])) in text

    assert f"{len(master):,}" + " recordings" in text
