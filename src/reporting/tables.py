"""The table engine -- one builder, four writers, one provenance record (Phase 85).

Thirty numbered tables (T01-T30) end up in a thesis, a Q1 paper and a dashboard.
Written by hand, or by thirty small scripts, they drift: T02 rounds a share to
two places and T11 to four, one table calls it *sensitivity* and another
*recall*, and a value copied into the DOCX stops matching the CSV it came from
the first time a run is repeated. This module exists so that a table is
**declared once** and rendered four ways from the same numbers.

The four writers and what each is for
-------------------------------------
``csv``
    **The source of truth (T85.2).** Full numeric precision, no rounding, no
    formatting, no provenance columns mixed in with the data. Everything else
    is a rendering of this file, and the T85.7 gate proves it by reading the
    rendered forms back and comparing them cell by cell.
``docx``
    Publication styling for the thesis (T85.3): a bold numbered caption, a grid
    table, and a small-print provenance line naming the experiment and the
    source file.
``latex``
    ``booktabs`` float for direct paste into the paper (T85.4), with the
    provenance carried in ``%`` comments so it travels with the file without
    appearing in the printed caption.
``md``
    The plain-text rendering used for README fragments and for eyeballing a
    table in a terminal or a diff.

Rounding (T85.6) is a property of the COLUMN, not of the writer
---------------------------------------------------------------
Metrics render at 3 decimals, percentages at 1, counts as integers with
thousands separators, and anything already formatted upstream (``0.8588 +/-
0.0255``) passes through untouched. Declaring the kind per column is why the
same value cannot be 0.86 in the DOCX and 0.8588 in the LaTeX.

Kinds are **declared** in the :class:`TableSpec` wherever the meaning is not
obvious from the name. :func:`infer_kind` exists as a fallback for wide source
files, and whatever it decided is written into the ``.meta.json`` -- an inferred
rounding rule that nobody can audit afterwards is the kind of quiet
transformation rule 1 is about.

Provenance (T85.5)
------------------
Every table writes ``<stem>.meta.json`` beside it holding the table id, title,
caption, experiment id, the resolved column kinds, and for each source CSV its
path, size, mtime and **sha256**. The digest is the part that matters: results
directories are rewritten as experiments are re-run, and a table whose numbers
came from a superseded CSV is indistinguishable from a current one without it.

``exp_id`` is never left blank. A setup table that is not bound to an experiment
records ``n/a (not experiment-bound)`` explicitly, so a missing value and a
genuinely-absent one do not look alike.
"""

from __future__ import annotations

import hashlib
import math
import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from src.utils.io import ensure_dir, save_csv, save_json
from src.utils.logging_setup import get_logger

__all__ = [
    "Column",
    "TableSpec",
    "Table",
    "ColumnKind",
    "PLACES",
    "NOT_EXPERIMENT_BOUND",
    "NA_TEXT",
    "WRITERS",
    "infer_kind",
    "format_value",
    "formatted_frame",
    "build_table",
    "write_table",
    "write_csv",
    "write_markdown",
    "write_docx",
    "write_latex",
    "read_docx_table",
    "read_latex_table",
    "read_markdown_table",
    "source_fingerprint",
]

log = get_logger("reporting.tables")

ColumnKind = str

#: Decimal places per column kind (T85.6). ``None`` means "render as text".
PLACES: dict[str, int | None] = {
    "metric": 3,
    "percent": 1,
    "count": 0,
    "integer": 0,
    "seconds": 2,
    "text": None,
    "preformatted": None,
}

NOT_EXPERIMENT_BOUND = "n/a (not experiment-bound)"

#: What a NaN renders as in the display writers. Never ``0``, never blank --
#: a blank cell reads as "we did not report it" and a zero reads as a result.
NA_TEXT = "n/a"

WRITERS: tuple[str, ...] = ("csv", "md", "docx", "latex")

# Column-name fragments that identify a metric. Matched on the lowercased name
# at word-ish boundaries first, so ``n_records`` does not match ``recall``.
_METRIC_TOKENS = (
    "sensitivity", "specificity", "recall", "precision", "f1", "fbeta",
    "accuracy", "auc", "auroc", "auprc", "mcc", "brier", "ece", "kappa",
    "score", "ratio", "share", "fraction", "correlation", "importance",
    "mean", "sd", "std", "median", "iqr", "p25", "p75",
)
_PERCENT_TOKENS = ("_pct", "pct_", "percent", "percentage")
_COUNT_TOKENS = ("count", "total", "files", "records", "subjects", "folds", "classes")
_SECONDS_TOKENS = ("seconds", "_sec", "wall_time", "elapsed", "hours")


def infer_kind(name: str, values: Any = None) -> str:
    """Best guess at a column's kind from its name, with its dtype as a tiebreak.

    Deliberately conservative: anything it is not confident about becomes
    ``text``, which renders the value unchanged. A wrong guess that formats a
    string is harmless; a wrong guess that rounds an identifier is not.
    """
    import pandas as pd

    lowered = name.lower()

    if any(token in lowered for token in _PERCENT_TOKENS):
        return "percent"
    if any(token in lowered for token in _SECONDS_TOKENS):
        return "seconds"

    series = None
    if values is not None:
        series = pd.Series(values)
        if not pd.api.types.is_numeric_dtype(series):
            sample = series.dropna().astype(str)
            # "0.8588 +/- 0.0255" and friends were rendered upstream already.
            if len(sample) and sample.str.contains(r"\+/-").any():
                return "preformatted"
            return "text"

    if any(re.search(r"(^|_)" + re.escape(token) + r"($|_)", lowered)
           for token in _METRIC_TOKENS):
        return "metric"
    if any(token in lowered for token in _COUNT_TOKENS):
        return "count"
    if series is not None and pd.api.types.is_integer_dtype(series):
        return "count"
    if any(token in lowered for token in _METRIC_TOKENS):
        return "metric"
    return "text"


@dataclass(frozen=True)
class Column:
    """One column: where it comes from, what it is called, how it rounds."""

    name: str
    header: str | None = None
    kind: str | None = None
    places: int | None = None

    def label(self) -> str:
        return self.header if self.header is not None else self.name


@dataclass(frozen=True)
class TableSpec:
    """Everything about a table that is not its numbers."""

    table_id: str
    title: str
    caption: str
    sources: tuple[str, ...]
    columns: tuple[Column, ...] = ()
    exp_id: str = ""
    objective: str = ""
    dataset: str = ""
    notes: tuple[str, ...] = ()
    command: str = ""

    def slug(self) -> str:
        cleaned = re.sub(r"[^a-z0-9]+", "_", self.title.lower()).strip("_")
        return self.table_id + "_" + cleaned

    def experiment_id(self) -> str:
        return self.exp_id or NOT_EXPERIMENT_BOUND


@dataclass
class Table:
    """A :class:`TableSpec` bound to the frame that satisfies it."""

    spec: TableSpec
    frame: Any
    kinds: dict[str, str] = field(default_factory=dict)

    def headers(self) -> list[str]:
        declared = {c.name: c.label() for c in self.spec.columns}
        return [declared.get(str(c), str(c)) for c in self.frame.columns]


# ---------------------------------------------------------------------------
# formatting
# ---------------------------------------------------------------------------


def format_value(value: Any, kind: str, places: int | None = None) -> str:
    """Render one cell under the T85.6 rules. The ONLY place rounding happens."""
    import pandas as pd

    if value is None:
        return NA_TEXT
    try:
        if bool(pd.isna(value)):
            return NA_TEXT
    except (TypeError, ValueError):
        pass

    resolved = PLACES.get(kind) if places is None else places
    if resolved is None:
        return str(value)

    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isinf(number):
        return "inf" if number > 0 else "-inf"

    if kind in ("count", "integer"):
        return format(round(number), ",d")
    return format(number, "." + str(resolved) + "f")


def formatted_frame(table: Table) -> Any:
    """The whole table as display strings, headers applied. Writers use only this."""
    import pandas as pd

    rendered: dict[str, list[str]] = {}
    for column, header in zip(table.frame.columns, table.headers(), strict=True):
        name = str(column)
        declared = next((c for c in table.spec.columns if c.name == name), None)
        kind = table.kinds.get(name) or (declared.kind if declared else None)
        if kind is None:
            kind = infer_kind(name, table.frame[column])
            table.kinds[name] = kind
        places = declared.places if declared else None
        rendered[header] = [
            format_value(v, kind, places) for v in table.frame[column].tolist()
        ]
    return pd.DataFrame(rendered)


# ---------------------------------------------------------------------------
# building
# ---------------------------------------------------------------------------


def source_fingerprint(path: str | Path) -> dict[str, Any]:
    """Path, size, mtime and sha256 of one source file.

    The digest is what makes a stale table detectable: results CSVs get
    rewritten when an experiment is re-run, and without a content hash a table
    built from the superseded file looks exactly like one built from the
    current one.
    """
    candidate = Path(path)
    if not candidate.is_file():
        return {
            "path": str(candidate).replace("\\", "/"),
            "exists": False,
            "bytes": None,
            "sha256": None,
            "modified_utc": None,
        }
    digest = hashlib.sha256(candidate.read_bytes()).hexdigest()
    stat = candidate.stat()
    return {
        "path": str(candidate).replace("\\", "/"),
        "exists": True,
        "bytes": stat.st_size,
        "sha256": digest,
        "modified_utc": datetime.fromtimestamp(stat.st_mtime, tz=UTC).isoformat(),
    }


def build_table(spec: TableSpec, frame: Any) -> Table:
    """Bind a frame to a spec, selecting and ordering the declared columns."""
    import pandas as pd

    data = pd.DataFrame(frame)
    if spec.columns:
        missing = [c.name for c in spec.columns if c.name not in data.columns]
        if missing:
            raise KeyError(
                spec.table_id + ": source frame is missing declared column(s) "
                + ", ".join(missing)
                + ". Present: " + ", ".join(str(c) for c in data.columns)
            )
        data = data[[c.name for c in spec.columns]]
    if data.empty:
        raise ValueError(spec.table_id + ": refusing to write an empty table")

    kinds: dict[str, str] = {}
    for column in data.columns:
        name = str(column)
        declared = next((c for c in spec.columns if c.name == name), None)
        kinds[name] = (
            declared.kind if declared and declared.kind
            else infer_kind(name, data[column])
        )
    return Table(spec=spec, frame=data.reset_index(drop=True), kinds=kinds)


# ---------------------------------------------------------------------------
# writers
# ---------------------------------------------------------------------------


def _provenance_lines(table: Table) -> list[str]:
    spec = table.spec
    lines = [
        "Table " + spec.table_id + " -- " + spec.title,
        "Experiment: " + spec.experiment_id(),
    ]
    if spec.objective:
        lines.append("Objective: " + spec.objective)
    lines.append(
        "Source: " + ("; ".join(spec.sources) if spec.sources else "(none recorded)")
    )
    lines.append("Generated by PV-MEPCG / PulseVision at " + datetime.now(UTC).isoformat())
    return lines


def write_csv(table: Table, out_dir: str | Path) -> Path:
    """Full precision, data only. The source of truth (T85.2)."""
    target = ensure_dir(out_dir) / (table.spec.slug() + ".csv")
    return save_csv(table.frame, target)


def write_markdown(table: Table, out_dir: str | Path) -> Path:
    display = formatted_frame(table)
    headers = [str(c) for c in display.columns]
    lines = [
        "### " + table.spec.table_id + " - " + table.spec.title,
        "",
        table.spec.caption,
        "",
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * len(headers)) + "|",
    ]
    for row in display.itertuples(index=False):
        lines.append("| " + " | ".join(str(v) for v in row) + " |")
    for note in table.spec.notes:
        lines.extend(["", "> " + note])
    lines.append("")
    lines.extend("_" + line + "_" for line in _provenance_lines(table))
    target = ensure_dir(out_dir) / (table.spec.slug() + ".md")
    target.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return target


def write_docx(table: Table, out_dir: str | Path) -> Path:
    """Publication-styled DOCX with a caption and a provenance line (T85.3)."""
    from docx import Document
    from docx.shared import Pt

    display = formatted_frame(table)
    document = Document()

    heading = document.add_paragraph()
    run = heading.add_run(table.spec.table_id + ". " + table.spec.title)
    run.bold = True
    run.font.size = Pt(11)

    caption = document.add_paragraph(table.spec.caption)
    caption.runs[0].font.size = Pt(9)

    grid = document.add_table(rows=1, cols=len(display.columns))
    grid.style = "Table Grid"
    for cell, header in zip(grid.rows[0].cells, display.columns, strict=True):
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for cell_run in paragraph.runs:
                cell_run.bold = True
                cell_run.font.size = Pt(9)
    for row in display.itertuples(index=False):
        cells = grid.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for cell_run in paragraph.runs:
                    cell_run.font.size = Pt(9)

    for note in table.spec.notes:
        note_paragraph = document.add_paragraph(note)
        note_paragraph.runs[0].font.size = Pt(8)
        note_paragraph.runs[0].italic = True

    for line in _provenance_lines(table):
        provenance = document.add_paragraph(line)
        provenance.runs[0].font.size = Pt(7)
        provenance.runs[0].italic = True

    target = ensure_dir(out_dir) / (table.spec.slug() + ".docx")
    document.save(str(target))
    return target


_LATEX_ESCAPES: tuple[tuple[str, str], ...] = (
    ("\\", "\\textbackslash{}"),
    ("&", "\\&"),
    ("%", "\\%"),
    ("$", "\\$"),
    ("#", "\\#"),
    ("_", "\\_"),
    ("{", "\\{"),
    ("}", "\\}"),
    ("~", "\\textasciitilde{}"),
    ("^", "\\textasciicircum{}"),
)


def _latex_escape(text: Any) -> str:
    out = str(text)
    # Backslash first, or the escapes introduced below would be escaped again.
    out = out.replace("\\", "\x00")
    for plain, escaped in _LATEX_ESCAPES[1:]:
        out = out.replace(plain, escaped)
    return out.replace("\x00", "\\textbackslash{}")


def _latex_unescape(text: str) -> str:
    out = str(text)
    for plain, escaped in reversed(_LATEX_ESCAPES[1:]):
        out = out.replace(escaped, plain)
    return out.replace("\\textbackslash{}", "\\")


def write_latex(table: Table, out_dir: str | Path) -> Path:
    """A ``booktabs`` float, provenance in ``%`` comments (T85.4).

    The provenance is a comment rather than caption text on purpose: it must
    travel with the file so the source is never lost, but a paper caption that
    reads "Source: outputs/01_dataset_audit/..." is not publishable.
    """
    display = formatted_frame(table)
    alignment = "l" + "r" * (len(display.columns) - 1)

    lines = ["% " + line for line in _provenance_lines(table)]
    lines += [
        "\\begin{table}[htbp]",
        "  \\centering",
        "  \\caption{" + _latex_escape(table.spec.caption) + "}",
        "  \\label{tab:" + table.spec.table_id.lower() + "}",
        "  \\begin{tabular}{" + alignment + "}",
        "    \\toprule",
        "    " + " & ".join(_latex_escape(c) for c in display.columns) + " \\\\",
        "    \\midrule",
    ]
    for row in display.itertuples(index=False):
        lines.append("    " + " & ".join(_latex_escape(v) for v in row) + " \\\\")
    lines += ["    \\bottomrule", "  \\end{tabular}"]
    for note in table.spec.notes:
        lines.append("  \\par\\footnotesize " + _latex_escape(note))
    lines += ["\\end{table}", ""]

    target = ensure_dir(out_dir) / (table.spec.slug() + ".tex")
    target.write_text("\n".join(lines), encoding="utf-8")
    return target


def _write_meta(table: Table, out_dir: str | Path, written: dict[str, Path]) -> Path:
    from src.utils.run_manifest import git_info

    spec = table.spec
    meta = {
        "table_id": spec.table_id,
        "title": spec.title,
        "caption": spec.caption,
        "exp_id": spec.experiment_id(),
        "objective": spec.objective or None,
        "dataset": spec.dataset or None,
        "framework": "PV-MEPCG / PulseVision",
        "generated_utc": datetime.now(UTC).isoformat(),
        "git": git_info(),
        "command": spec.command or None,
        "rounding_rules": dict(PLACES),
        "column_kinds": dict(table.kinds),
        "n_rows": len(table.frame),
        "sources": [source_fingerprint(s) for s in spec.sources],
        "written": {fmt: str(path).replace("\\", "/") for fmt, path in written.items()},
        "notes": list(spec.notes),
    }
    target = ensure_dir(out_dir) / (spec.slug() + ".meta.json")
    return save_json(meta, target)


def write_table(
    table: Table,
    out_dir: str | Path,
    *,
    formats: tuple[str, ...] = WRITERS,
    evidence_index: str | Path | None = None,
) -> dict[str, Path]:
    """Render one table through every requested writer plus its ``.meta.json``."""
    from src.utils.evidence import register_evidence

    dispatch = {
        "csv": write_csv,
        "md": write_markdown,
        "docx": write_docx,
        "latex": write_latex,
    }
    unknown = [f for f in formats if f not in dispatch]
    if unknown:
        raise ValueError("unknown table writer(s): " + ", ".join(unknown))

    written: dict[str, Path] = {fmt: dispatch[fmt](table, out_dir) for fmt in formats}
    written["meta"] = _write_meta(table, out_dir, written)

    register_evidence(
        table.spec.table_id,
        written.get("csv", next(iter(written.values()))),
        metric_or_asset=table.spec.title,
        objective=table.spec.objective,
        experiment_id=table.spec.exp_id,
        dataset=table.spec.dataset,
        source_data="; ".join(table.spec.sources),
        command=table.spec.command,
        index_path=evidence_index,
    )
    log.info(
        "%s: wrote %s",
        table.spec.table_id,
        ", ".join(p.name for p in written.values()),
    )
    return written


# ---------------------------------------------------------------------------
# readers -- the T85.7 gate has to read the rendered forms BACK
# ---------------------------------------------------------------------------


def read_markdown_table(path: str | Path) -> list[list[str]]:
    """Header row plus body rows from a written Markdown table."""
    rows: list[list[str]] = []
    for line in Path(path).read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        cells = [c.strip() for c in stripped[1:-1].split("|")]
        if all(c and set(c) <= set("-: ") for c in cells):
            continue  # the |---|---| separator
        rows.append(cells)
    return rows


def read_docx_table(path: str | Path, index: int = 0) -> list[list[str]]:
    """Header row plus body rows from a written DOCX table."""
    from docx import Document

    grid = Document(str(path)).tables[index]
    return [[cell.text.strip() for cell in row.cells] for row in grid.rows]


def read_latex_table(path: str | Path) -> list[list[str]]:
    """Header row plus body rows from a written LaTeX table."""
    rows: list[list[str]] = []
    inside = False
    for raw in Path(path).read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line.startswith("\\begin{tabular}"):
            inside = True
            continue
        if line.startswith("\\end{tabular}"):
            break
        if not inside or not line.endswith("\\\\"):
            continue
        body = line[:-2].strip()
        # Split on column separators only: an escaped ``\&`` is cell content.
        cells = re.split(r"(?<!\\)&", body)
        rows.append([_latex_unescape(cell.strip()) for cell in cells])
    return rows
