"""Per-recording and per-experiment reports (Phase 107).

Three deliverables, one rule between them: **a report may only restate what a
run produced.** Nothing here computes a metric. The per-recording report renders
a `PredictionResult` from `src.inference.predictor`; the experiment report reads
`per_fold_metrics.csv` and `aggregate_metrics.csv`; the batch export writes the
prediction objects it is handed. If a number is not in one of those, it does not
appear.

## DOCX rather than PDF

T107.1 allows either. DOCX is chosen because `python-docx` is already pinned and
installed, whereas a PDF writer would add a dependency for a format nobody in
this project reads programmatically — and the T107.7 gate has to *check* the
disclaimer is present, which means reading the file back. A DOCX round-trips
through `python-docx`; a PDF would need a second parser to verify what the first
one wrote.

## Top contributing features, and when there are none

T107.2 asks for the features that drove *this* prediction. For a linear model
that is well defined and cheap: the contribution of feature *i* is
`coef[i] * z[i]`, where `z` is the value after the pipeline's imputer and
scaler. Those terms sum, with the intercept, to the log-odds — so the report can
show a decomposition that adds up rather than a ranking of unrelated numbers.

For a forest or a boosted ensemble it is **not** available from the fitted model
alone; it needs SHAP, which is Part VIII's explainability work. The report says
so in those words rather than substituting global feature importance, which
answers a different question — "what matters across the corpus" is not "what
drove this recording".

## The stamp

Every report carries the screening disclaimer and the model version block:
model id, estimator, when it was saved, how many records it was fitted on, and
the package versions. A report without those is a number in a document, and a
number in a document outlives the run that produced it.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import numpy as np

from src.inference.predictor import DISCLAIMER, PredictionResult
from src.utils.logging_setup import get_logger

__all__ = [
    "CONTRIBUTION_UNAVAILABLE",
    "REPORTED_METRICS",
    "FeatureContribution",
    "ReportError",
    "batch_export",
    "feature_contributions",
    "render_experiment_report",
    "render_sample_report",
    "report_for_recording",
]

log = get_logger("reporting.sample_report")

CONTRIBUTION_UNAVAILABLE = (
    "Per-recording feature contributions are not available for this estimator. "
    "They are exact for a linear model, where each term is a coefficient times "
    "the scaled feature value and the terms sum to the log-odds. For a forest or "
    "a boosted ensemble they require SHAP, which is Part VIII's explainability "
    "work. Global feature importance is deliberately NOT substituted here: what "
    "matters across the corpus is a different question from what drove this "
    "recording."
)


class ReportError(RuntimeError):
    """The report cannot be assembled from what was provided."""


@dataclass(frozen=True)
class FeatureContribution:
    """One feature's signed contribution to this recording's log-odds."""

    name: str
    #: The raw feature value, as extracted.
    value: float
    #: The value after the pipeline's imputer and scaler.
    scaled: float
    coefficient: float
    #: ``coefficient * scaled``. Positive pushes toward the positive class.
    contribution: float

    @property
    def direction(self) -> str:
        return "toward abnormal" if self.contribution > 0 else "toward normal"


# ---------------------------------------------------------------------------
# T107.2 -- what drove this prediction
# ---------------------------------------------------------------------------


def feature_contributions(
    bundle: Any, vector: np.ndarray, *, top: int = 10
) -> tuple[list[FeatureContribution], str | None]:
    """``(contributions, unavailable_reason)`` for one recording.

    Returns an empty list and a reason when the estimator is not linear, rather
    than raising: a report for a forest is still a useful report, it simply
    cannot carry this section.
    """
    from sklearn.pipeline import Pipeline

    pipeline = bundle.pipeline
    if not isinstance(pipeline, Pipeline):
        return [], CONTRIBUTION_UNAVAILABLE

    estimator = pipeline.steps[-1][1]
    coefficients = getattr(estimator, "coef_", None)
    if coefficients is None:
        return [], CONTRIBUTION_UNAVAILABLE

    coef = np.asarray(coefficients, dtype=np.float64)
    if coef.ndim == 2:
        if coef.shape[0] != 1:
            # Multiclass linear models have one coefficient vector per class;
            # a single "contribution" is not defined without naming the class.
            return [], CONTRIBUTION_UNAVAILABLE
        coef = coef[0]

    row = vector.reshape(1, -1)
    for name, step in pipeline.steps[:-1]:
        try:
            row = step.transform(row)
        except Exception as error:
            # Broad on purpose: any step can fail for its own reasons, and the
            # thing that makes the failure actionable is which step it was.
            raise ReportError(
                "the pipeline step " + repr(name) + " could not transform this "
                "recording's features: " + str(error)
            ) from error

    scaled = np.asarray(row, dtype=np.float64).ravel()
    if scaled.size != coef.size:
        # A selector dropped columns; the surviving names are unknown here.
        return [], CONTRIBUTION_UNAVAILABLE

    names = list(bundle.feature_names)
    terms = coef * scaled
    order = np.argsort(np.abs(terms))[::-1][:top]

    return [
        FeatureContribution(
            name=names[int(index)],
            value=float(vector[int(index)]),
            scaled=float(scaled[int(index)]),
            coefficient=float(coef[int(index)]),
            contribution=float(terms[int(index)]),
        )
        for index in order
    ], None


# ---------------------------------------------------------------------------
# figures
# ---------------------------------------------------------------------------


def _waveform_and_spectrogram(signal: np.ndarray, fs: int, out_dir: Path, stem: str) -> list[Path]:
    """The two figures T107.2 names, at the project's 300 dpi print settings."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from src.reporting.plot_style import SEQUENTIAL_CMAP, figure, save_figure

    seconds = np.arange(signal.size, dtype=np.float64) / float(fs)

    fig, ax = figure("double")
    ax.plot(seconds, signal, linewidth=0.6)
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("Amplitude (normalized)")
    ax.set_title("Preprocessed waveform")
    ax.margins(x=0)
    waveform = save_figure(
        fig,
        out_dir / (stem + "_waveform.png"),
        source="Preprocessed signal at " + str(fs) + " Hz, after the shared pipeline.",
    )

    fig, ax = figure("double")
    ax.specgram(signal, Fs=fs, NFFT=512, noverlap=256, cmap=SEQUENTIAL_CMAP, scale="dB")
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("Frequency (Hz)")
    ax.set_title("Spectrogram")
    ax.set_ylim(0, min(500, fs / 2))
    spectrogram = save_figure(
        fig,
        out_dir / (stem + "_spectrogram.png"),
        source="Short-time Fourier transform, 512-sample window, 50% overlap.",
    )
    plt.close("all")
    return [waveform, spectrogram]


# ---------------------------------------------------------------------------
# T107.1 / T107.3 -- the per-recording report
# ---------------------------------------------------------------------------


def render_sample_report(
    result: PredictionResult,
    out_path: str | Path,
    *,
    bundle: Any | None = None,
    vector: np.ndarray | None = None,
    signal: np.ndarray | None = None,
    fs: int | None = None,
    figures_dir: str | Path | None = None,
) -> Path:
    """One recording's report as DOCX. Returns the path written.

    `bundle` and `vector` are optional: without them the contributions section
    states that it is unavailable, which is the honest rendering for a report
    generated from a stored result rather than a live prediction.
    """
    from docx import Document
    from docx.shared import Inches, Pt

    from src.utils.io import ensure_dir

    target = Path(out_path)
    ensure_dir(target.parent)
    stem = target.stem
    images_dir = Path(figures_dir) if figures_dir is not None else target.parent
    ensure_dir(images_dir)

    document = Document()
    document.add_heading("PV-MEPCG / PulseVision — recording report", level=1)

    # T107.3: the disclaimer is the first thing in the document, before any
    # number. A disclaimer at the end is a disclaimer read after the decision.
    scope = document.add_paragraph()
    run = scope.add_run(DISCLAIMER)
    run.bold = True
    run.font.size = Pt(10)

    document.add_heading("Recording", level=2)
    _key_values(
        document,
        [
            ("Source", result.source),
            ("Task", result.task),
            ("Duration", _seconds(result.quality.get("duration_seconds"))),
            ("Original sample rate", _hz(result.quality.get("original_sample_rate_hz"))),
            ("Channels", str(result.quality.get("channels", "n/a"))),
        ],
    )

    document.add_heading("Screening result", level=2)
    _key_values(
        document,
        [
            ("Predicted class", result.predicted_class),
            ("Confidence", format(result.confidence, ".4f")),
            ("Margin over the runner-up", format(result.margin, ".4f")),
            (
                "Low confidence",
                "YES — the top two classes are within "
                + format(result.low_confidence_margin, ".2f")
                if result.low_confidence
                else "no",
            ),
            (
                "Operating point",
                "argmax at " + format(result.operating_threshold, ".2f")
                if result.operating_threshold is not None
                else "argmax",
            ),
        ],
    )
    document.add_paragraph(result.operating_point_note)

    document.add_heading("Class probabilities", level=3)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Class"
    header[1].text = "Probability"
    for name, value in result.probabilities.items():
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = format(value, ".4f")

    if signal is not None and fs is not None:
        document.add_heading("Signal", level=2)
        for image in _waveform_and_spectrogram(signal, int(fs), images_dir, stem):
            document.add_picture(str(image), width=Inches(6.0))

    document.add_heading("What drove this prediction", level=2)
    contributions: list[FeatureContribution] = []
    reason: str | None = CONTRIBUTION_UNAVAILABLE
    if bundle is not None and vector is not None:
        contributions, reason = feature_contributions(bundle, vector)
    if contributions:
        document.add_paragraph(
            "Each term is the model coefficient times the scaled feature value. "
            "The terms and the intercept sum to the log-odds, so this is a "
            "decomposition of the decision rather than a ranking of unrelated "
            "numbers."
        )
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        head = table.rows[0].cells
        # "Scaled" is in the table because without it the arithmetic does not
        # visibly check out: the contribution is coefficient x *scaled*, not
        # coefficient x raw value, and a reader multiplying the two printed
        # numbers would get a different answer and distrust the table.
        for index, label in enumerate(
            ("Feature", "Value", "Scaled", "Coefficient", "Contribution")
        ):
            head[index].text = label
        for item in contributions:
            cells = table.add_row().cells
            cells[0].text = item.name
            cells[1].text = format(item.value, ".4f")
            cells[2].text = format(item.scaled, ".4f")
            cells[3].text = format(item.coefficient, ".4f")
            cells[4].text = format(item.contribution, "+.4f") + " (" + item.direction + ")"
    else:
        document.add_paragraph(reason or CONTRIBUTION_UNAVAILABLE)

    if result.warnings:
        document.add_heading("Warnings", level=2)
        for warning in result.warnings:
            document.add_paragraph(warning, style="List Bullet")

    # T107.3: the model version block.
    document.add_heading("Model version", level=2)
    model = result.model
    _key_values(
        document,
        [
            ("Model id", str(model.get("model_id", "n/a"))),
            ("Estimator", str(model.get("estimator_class", "n/a"))),
            ("Features", str(model.get("n_features", "n/a"))),
            ("Saved at", str(model.get("saved_at", "n/a"))),
            ("Records fitted", str(model.get("n_records_fitted", "n/a"))),
            ("Selection rule", ", ".join(model.get("selection_rule") or []) or "n/a"),
            ("Bundle path", str(model.get("path", "n/a"))),
        ],
    )
    note = model.get("note")
    if note:
        document.add_paragraph(str(note))
    versions = model.get("package_versions") or {}
    if versions:
        document.add_paragraph(
            "Package versions: "
            + ", ".join(name + " " + str(value) for name, value in sorted(versions.items()))
        )

    document.add_paragraph(DISCLAIMER)
    document.save(str(target))
    log.info("wrote sample report -> %s", target)
    return target


def report_for_recording(
    path: str | Path,
    out_path: str | Path,
    *,
    task: str = "binary",
    record_uid: str | None = None,
    figures_dir: str | Path | None = None,
) -> tuple[Path, PredictionResult]:
    """One WAV to one finished report, in a single pass over the audio.

    The whole T107.1/T107.2 deliverable is this function: it scores the
    recording through `predict_recording` -- the only path from a WAV to a class
    -- and renders the **same** signal and the **same** feature vector that
    produced the number. Preprocessing and extraction run once. A caller that
    re-ran them to get pictures would be drawing a second computation and
    trusting it to agree.
    """
    from src.inference.predictor import predict_recording

    result, detail = predict_recording(path, task=task, record_uid=record_uid, with_detail=True)
    written = render_sample_report(
        result,
        out_path,
        bundle=detail.bundle,
        vector=detail.vector,
        signal=detail.signal,
        fs=detail.fs,
        figures_dir=figures_dir,
    )
    return written, result


def _key_values(document: Any, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _seconds(value: Any) -> str:
    return "n/a" if value is None else format(float(value), ".3f") + " s"


def _hz(value: Any) -> str:
    return "n/a" if value is None else str(int(value)) + " Hz"


# ---------------------------------------------------------------------------
# T107.4 -- the experiment-level report
# ---------------------------------------------------------------------------


def render_experiment_report(
    exp_dir: str | Path, out_path: str | Path, *, title: str | None = None
) -> Path:
    """Summarise one experiment run from the files it wrote.

    Reads `aggregate_metrics.csv`, `per_fold_metrics.csv`, `run_manifest.json`
    and `config_snapshot.yaml` — and computes nothing. Every number in the
    document is a cell from one of those files, formatted.
    """
    import pandas as pd
    from docx import Document
    from docx.shared import Pt

    from src.utils.io import ensure_dir

    source = Path(exp_dir)
    aggregate_path = source / "aggregate_metrics.csv"
    per_fold_path = source / "per_fold_metrics.csv"
    if not aggregate_path.is_file():
        raise ReportError("no aggregate_metrics.csv in " + str(source))

    aggregate = pd.read_csv(aggregate_path)
    manifest: dict[str, Any] = {}
    manifest_path = source / "run_manifest.json"
    if manifest_path.is_file():
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

    target = Path(out_path)
    ensure_dir(target.parent)

    document = Document()
    document.add_heading("PV-MEPCG / PulseVision — experiment report", level=1)
    run = document.add_paragraph().add_run(DISCLAIMER)
    run.bold = True
    run.font.size = Pt(10)

    document.add_heading(title or source.name, level=2)
    git = manifest.get("git") or {}
    _key_values(
        document,
        [
            ("Experiment directory", source.name),
            ("Run id", str(manifest.get("run_id", "n/a"))),
            ("Seed", str(manifest.get("seed", "n/a"))),
            ("Git commit", str(git.get("commit", "n/a"))),
            ("Started", str(manifest.get("started_utc", "n/a"))),
            ("Finished", str(manifest.get("finished_utc", "n/a"))),
        ],
    )

    document.add_heading("Aggregate metrics", level=2)
    document.add_paragraph(
        "Mean +/- SD across folds, read from aggregate_metrics.csv and rounded "
        "to three places by the shared table rules. Sensitivity and balanced "
        "accuracy are the selection metrics; accuracy alone is never the "
        "headline. Rows are the metrics research rule 6 requires, in that order."
    )
    _aggregate_table(document, aggregate)

    if per_fold_path.is_file():
        per_fold = pd.read_csv(per_fold_path)
        document.add_heading("Folds", level=2)
        _key_values(
            document,
            [
                ("Rows", str(len(per_fold))),
                ("Models", ", ".join(sorted(per_fold["model_id"].astype(str).unique()))),
                ("Folds per model", str(per_fold["fold_label"].nunique())),
            ],
        )

    # Provenance: which files this document was built from, and their digests.
    # Without it the report is a set of numbers whose sources cannot be checked
    # once the run directory moves on.
    from src.reporting.tables import source_fingerprint

    document.add_heading("Sources", level=2)
    fingerprints: list[tuple[str, str]] = []
    for candidate in (
        aggregate_path,
        per_fold_path,
        manifest_path,
        source / "config_snapshot.yaml",
    ):
        if not candidate.is_file():
            continue
        fingerprint = source_fingerprint(candidate)
        fingerprints.append((candidate.name, str(fingerprint["sha256"])[:16]))
    _key_values(document, [(name + " sha256", digest) for name, digest in fingerprints])

    document.add_paragraph(DISCLAIMER)
    document.save(str(target))
    log.info("wrote experiment report -> %s", target)
    return target


#: The metrics reported for every run, in the order research rule 6 states them.
#: A metric absent from a given `aggregate_metrics.csv` is skipped rather than
#: rendered as a zero -- EXP-A2 carries columns EXP-A1 does not, and vice versa.
REPORTED_METRICS: tuple[tuple[str, str], ...] = (
    ("sensitivity", "Sensitivity (recall, abnormal)"),
    ("specificity", "Specificity"),
    ("balanced_accuracy", "Balanced accuracy"),
    ("f1", "F1"),
    ("precision", "Precision"),
    ("roc_auc", "ROC AUC"),
    ("pr_auc", "PR AUC"),
    ("mcc", "MCC"),
    ("accuracy", "Accuracy"),
    ("brier", "Brier score"),
    ("ece", "Expected calibration error"),
)


def _aggregate_table(document: Any, frame: Any) -> None:
    """Metrics down the rows, models across the columns, formatted once.

    The raw frame is 71 columns wide (a mean, an SD and an n for every metric,
    plus fold sizes and confusion-matrix counts). Rendering it whole produces a
    table nobody can read, and rendering `str(value)` publishes a metric at
    seventeen significant figures. Both are fixed here: the mean and SD are
    paired into one cell and rounded by `tables.format_value`, which is the only
    place in this project rounding happens.
    """
    from src.reporting.tables import NA_TEXT, format_value

    models = [str(value) for value in frame["model_id"].tolist()]
    rows: list[tuple[str, list[str]]] = []
    for column, label in REPORTED_METRICS:
        mean_column = column + "_mean"
        if mean_column not in frame.columns:
            continue
        sd_column = column + "_sd"
        cells: list[str] = []
        for position in range(len(frame)):
            mean = format_value(frame[mean_column].iloc[position], "metric")
            if mean == NA_TEXT or sd_column not in frame.columns:
                cells.append(mean)
                continue
            cells.append(mean + " +/- " + format_value(frame[sd_column].iloc[position], "metric"))
        rows.append((label, cells))

    if not rows:
        raise ReportError("aggregate_metrics.csv carries none of the reported metrics")

    table = document.add_table(rows=1, cols=len(models) + 1)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Metric"
    for index, model in enumerate(models):
        header[index + 1].text = model
    for label, cells in rows:
        written = table.add_row().cells
        written[0].text = label
        for index, value in enumerate(cells):
            written[index + 1].text = value


# ---------------------------------------------------------------------------
# T107.5 -- batch export
# ---------------------------------------------------------------------------


def batch_export(
    results: list[PredictionResult], out_dir: str | Path, *, stem: str = "batch_predictions"
) -> dict[str, Path]:
    """Write a batch of predictions as both CSV and JSON.

    The CSV is one row per recording with the class probabilities widened into
    columns; the JSON keeps the full structure including timings, warnings and
    the model block. Both carry the disclaimer — the CSV as a leading comment
    line, since a spreadsheet has nowhere else to put one.
    """
    import pandas as pd

    from src.utils.io import ensure_dir

    directory = ensure_dir(out_dir)
    if not results:
        raise ReportError("no predictions to export")

    rows: list[dict[str, Any]] = []
    for result in results:
        row: dict[str, Any] = {
            "source": result.source,
            "task": result.task,
            "predicted_class": result.predicted_class,
            "confidence": result.confidence,
            "margin": result.margin,
            "low_confidence": result.low_confidence,
            "n_missing_features": result.n_missing_features,
            "model_id": result.model.get("model_id"),
            "seconds_total": result.timings_seconds.get("total"),
        }
        for name, value in result.probabilities.items():
            row["proba_" + name] = value
        rows.append(row)

    frame = pd.DataFrame(rows)
    csv_path = directory / (stem + ".csv")
    with csv_path.open("w", encoding="utf-8", newline="") as handle:
        handle.write("# " + DISCLAIMER + "\n")
        frame.to_csv(handle, index=False, lineterminator="\n")

    json_path = directory / (stem + ".json")
    payload = {
        "disclaimer": DISCLAIMER,
        "n_records": len(results),
        "predictions": [result.to_dict() for result in results],
    }
    # allow_nan=False on purpose: json.dumps otherwise emits bare `NaN`, which is
    # not JSON and which every strict parser rejects. A non-finite value has to
    # become null before it gets here, which is what `_finite` in the predictor
    # does -- this raises rather than shipping a file the dashboard cannot read.
    json_path.write_text(
        json.dumps(payload, indent=2, allow_nan=False), encoding="utf-8", newline="\n"
    )

    log.info("exported %d predictions -> %s, %s", len(results), csv_path, json_path)
    return {"csv": csv_path, "json": json_path}
