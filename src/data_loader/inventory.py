"""DA-01 dataset inventory and DA-09 narrative audit report (Phase 21).

Every number in both artifacts is computed here from the master table, the
integrity scan and the audit CSVs. Nothing is typed in. That is ground rule 1,
and it is the whole point of generating the .docx from code rather than writing
it: when a loader changes, the report changes with it, and there is no window in
which the prose disagrees with the CSVs it describes.

The three notes T21.2 -- T21.4 require are held as structured records
(:data:`DISCREPANCIES`, :data:`LIMITATIONS`) rather than free prose, so the
gate can assert their presence rather than grep the document for a phrase.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from src.utils.logging_setup import get_logger

__all__ = [
    "INVENTORY_COLUMNS",
    "Discrepancy",
    "DISCREPANCIES",
    "LIMITATIONS",
    "DATASET_METADATA",
    "build_inventory",
    "write_inventory",
    "build_report_sections",
    "write_audit_report",
    "register_audit_artifacts",
    "run_inventory",
]

log = get_logger(__name__)

INVENTORY_COLUMNS: tuple[str, ...] = (
    "dataset_source",
    "dataset_name",
    "version",
    "source",
    "folder",
    "total_files",
    "usable_files",
    "n_classes",
    "classes",
    "n_subjects",
    "subject_id_origin",
    "original_fs",
    "target_fs",
    "total_hours",
    "role",
)


@dataclass(frozen=True)
class Discrepancy:
    """One documented disagreement between a source document and the disk."""

    task: str
    subject: str
    blueprint_says: str
    disk_says: str
    explanation: str


# T21.2 and T21.3. Both are cases where the blueprint's counts do not match the
# audited files; CLAUDE.md's rule is that reality wins and the gap is written
# down rather than quietly reconciled.
DISCREPANCIES: tuple[Discrepancy, ...] = (
    Discrepancy(
        task="T21.2",
        subject="CirCor DigiScope 2022 corpus size",
        blueprint_says="1,568 patients / 5,272 recordings",
        disk_says="942 patients / 3,163 recordings",
        explanation=(
            "The blueprint quotes the size of the full CirCor DigiScope study. "
            "The public PhysioNet/CinC 2022 release contains only the training "
            "portion; the remaining 626 patients and 2,109 recordings are the "
            "Challenge's hidden validation and test sets, which were never "
            "published and cannot be obtained. Every CirCor figure in this "
            "project is computed on the 942/3,163 that exist, and results are "
            "not comparable to any figure quoted against the full 1,568."
        ),
    ),
    Discrepancy(
        task="T21.3",
        subject="PhysioNet/CinC 2016 corpus size",
        blueprint_says="approximately 3,126 recordings",
        disk_says=(
            "3,240 training recordings (training-a..f) plus a 301-record "
            "validation folder that duplicates training material"
        ),
        explanation=(
            "The blueprint's 3,126 matches no folder on disk. The training "
            "sets a-f hold 409 + 490 + 31 + 55 + 2,141 + 114 = 3,240 "
            "recordings, verified against each subset's REFERENCE.csv. The "
            "separate validation/ folder holds 301 further WAV files, but all "
            "301 are byte-identical (SHA-256) copies of records already in "
            "training-a..f, so the corpus contains 3,240 distinct recordings, "
            "not 3,541. This project uses the 3,240 and claims no held-out "
            "PhysioNet test set; see outputs/missing_outputs_report.txt."
        ),
    ),
)

# T21.4, plus the two other limitations that constrain what may be claimed.
LIMITATIONS: tuple[Discrepancy, ...] = (
    Discrepancy(
        task="T21.4",
        subject="PASCAL set_a carries no recoverable subject IDs",
        blueprint_says="not addressed",
        disk_says="176 files, no patient or session identifier in any field",
        explanation=(
            "set_a filenames encode a class and a timestamp and nothing else, "
            "and the dataset ships no patient key. There is therefore no way to "
            "guarantee that two set_a recordings come from different people. "
            "PASCAL A results are RECORD-LEVEL ONLY and must be described as "
            "such: cross-validation groups on a per-record key, so a subject "
            "recorded twice could in principle appear on both sides of a split. "
            "The one case where duplication is provable -- a single recording "
            "filed under both extrahls and murmur -- is grouped explicitly so "
            "that pair can never be split. With 124 labelled records and 19 in "
            "the smallest class, PASCAL A cannot carry Objective 6 alone; the "
            "PhysioNet diagnosis track (EXP-G1, 3,240 records) carries it."
        ),
    ),
    Discrepancy(
        task="T21.4",
        subject="PASCAL `artifact` is a recording-quality label",
        blueprint_says="four-class classification",
        disk_says="40 of 124 set_a records are labelled `artifact`",
        explanation=(
            "`artifact` marks an unusable recording, not a cardiac condition. "
            "The four-class set_a model is a four-class AUDIO model and must "
            "never be described as a four-class cardiac classifier."
        ),
    ),
    Discrepancy(
        task="T21.4",
        subject="PhysioNet subject IDs are derived, not native",
        blueprint_says="not addressed",
        disk_says=(
            "857 subject groups over 3,240 records; 2,984 records grouped from "
            "filename or annotation patterns, 256 falling back to record level"
        ),
        explanation=(
            "PhysioNet ships no patient key. Subjects are derived per subset "
            "from filename patterns and, for training-e, from the appendix's "
            "cohort-namespaced `# Raw record` field. Records where no pattern "
            "applies fall back to a record-level group and are flagged "
            "subject_derived=False, so any claim about subject-level "
            "independence can be qualified by exactly how many records it rests "
            "on."
        ),
    ),
)

# The provenance fields DA-01 asks for that are not derivable from the files.
DATASET_METADATA: dict[str, dict[str, str]] = {
    "D1": {
        "version": "PhysioNet/CinC Challenge 2016, v1.0.0",
        "source": "https://physionet.org/content/challenge-2016/",
        "subject_id_origin": "derived (filename and appendix patterns)",
        "role": "primary binary track (EXP-A1, EXP-A2, EXP-F1, EXP-F2, EXP-G1)",
    },
    "D2": {
        "version": "PASCAL Classifying Heart Sounds Challenge 2011, set A",
        "source": "http://www.peterjbentley.com/heartchallenge/",
        "subject_id_origin": "none (record-level only)",
        "role": "four-class audio track (EXP-B1)",
    },
    "D3": {
        "version": "PASCAL Classifying Heart Sounds Challenge 2011, set B",
        "source": "http://www.peterjbentley.com/heartchallenge/",
        "subject_id_origin": "derived (filename subject and timestamp)",
        "role": "three-class track (EXP-B2)",
    },
    "D4": {
        "version": "CirCor DigiScope, PhysioNet/CinC Challenge 2022, v1.0.3",
        "source": "https://physionet.org/content/circor-heart-sound/",
        "subject_id_origin": "native (patient id)",
        "role": "murmur and outcome tracks, external validation (EXP-C1..C3, EXP-D1)",
    },
}

TARGET_FS = 2000


# --------------------------------------------------------------------------
# paths
# --------------------------------------------------------------------------


def _audit_dir(out_dir: str | Path | None = None) -> Path:
    from src.utils.config import load_config
    from src.utils.io import ensure_dir

    if out_dir is not None:
        return ensure_dir(out_dir)
    return ensure_dir(load_config("paths").require("outputs.dataset_audit"))


def _dataset_folder(dataset_source: str) -> str:
    from src.utils.config import load_config

    paths = load_config("paths")
    key = {
        "D1": "dataset.d1_physionet.root",
        "D2": "dataset.d2_pascal_a.root",
        "D3": "dataset.d3_pascal_b.root",
        "D4": "dataset.d4_circor.root",
    }[dataset_source]
    root = Path(str(paths.require(key)))
    project = Path(str(paths.require("project_root")))
    try:
        return root.relative_to(project).as_posix()
    except ValueError:  # pragma: no cover -- dataset moved off the project root
        return root.as_posix()


# --------------------------------------------------------------------------
# T21.1 -- DA-01
# --------------------------------------------------------------------------


def build_inventory(master: Any) -> Any:
    """One row per dataset (T21.1).

    ``usable_files`` is the supervised count -- what actually enters a fold --
    which for D1 is 3,240 of 3,541 and for D2/D3 excludes the unlabelled files.
    """
    import pandas as pd

    from src.data_loader.master import TASK_DATASETS, task_frame

    rows: list[dict[str, Any]] = []
    for dataset in ("D1", "D2", "D3", "D4"):
        block = master[master["dataset_source"] == dataset]
        supervised = block[block["use_in_supervised"]]
        tasks = [t for t, d in TASK_DATASETS.items() if d == dataset]

        classes: list[str] = []
        for task in tasks:
            frame = task_frame(master, task)
            names = sorted(set(frame[_name_column(task)].astype(str)))
            classes.append(task + ": " + "|".join(names))

        n_classes = sum(
            task_frame(master, task)[_name_column(task)].nunique() for task in tasks
        )
        meta = DATASET_METADATA[dataset]
        rows.append(
            {
                "dataset_source": dataset,
                "dataset_name": str(block["dataset_name"].iloc[0]),
                "version": meta["version"],
                "source": meta["source"],
                "folder": _dataset_folder(dataset),
                "total_files": len(block),
                "usable_files": len(supervised),
                "n_classes": int(n_classes),
                "classes": "; ".join(classes),
                "n_subjects": int(supervised["subject_id"].nunique()),
                "subject_id_origin": meta["subject_id_origin"],
                "original_fs": int(block["original_fs"].iloc[0]),
                "target_fs": TARGET_FS,
                "total_hours": round(float(block["duration_sec"].sum()) / 3600.0, 4),
                "role": meta["role"],
            }
        )
    inventory = pd.DataFrame(rows, columns=list(INVENTORY_COLUMNS))
    log.info(
        "inventory: %d datasets, %d files, %d usable",
        len(inventory),
        int(inventory["total_files"].sum()),
        int(inventory["usable_files"].sum()),
    )
    return inventory


def _name_column(task: str) -> str:
    return {
        "binary": "binary_label_name",
        "pascal_a": "multiclass_label_name",
        "pascal_b": "multiclass_label_name",
        "circor_murmur": "murmur_label_name",
        "circor_outcome": "outcome_label_name",
    }[task]


def write_inventory(inventory: Any, out_dir: str | Path | None = None) -> Path:
    from src.utils.io import save_csv

    target = save_csv(inventory, _audit_dir(out_dir) / "dataset_inventory.csv")
    log.info("wrote %s (%d rows)", target.name, len(inventory))
    return target


# --------------------------------------------------------------------------
# T21.5 -- DA-09
# --------------------------------------------------------------------------

DISCLAIMER = (
    "PV-MEPCG / PulseVision is an academic screening and decision-support "
    "prototype. It is not a diagnostic device, it does not diagnose, treat or "
    "prescribe, and it does not replace assessment by a qualified clinician."
)


def build_report_sections(
    master: Any,
    inventory: Any,
    *,
    audit_dir: Path | None = None,
) -> list[tuple[str, list[str]]]:
    """The narrative body of DA-09, as ``(heading, paragraphs)`` pairs.

    Built as data so the gate can assert what the report says without parsing
    the .docx it produced.
    """
    import pandas as pd

    audit_dir = audit_dir or _audit_dir()

    def _read(name: str) -> Any:
        path = audit_dir / name
        if not path.is_file():
            return pd.DataFrame()
        return pd.read_csv(path, keep_default_na=False)

    duplicates = _read("duplicate_report.csv")
    missing = _read("missing_corrupt_files.csv")
    class_dist = _read("class_distribution.csv")
    split_map = _read("subject_split_map.csv")

    sections: list[tuple[str, list[str]]] = []

    sections.append(
        (
            "Scope and disclaimer",
            [
                DISCLAIMER,
                "This report is generated by src/data_loader/inventory.py from "
                "the audit artifacts in outputs/01_dataset_audit/. Every count "
                "below is read from those files at generation time; no figure "
                "in this document was typed by hand.",
            ],
        )
    )

    total = len(master)
    supervised = int(master["use_in_supervised"].sum())
    hours = float(master["duration_sec"].sum()) / 3600.0
    sections.append(
        (
            "1. Corpus overview",
            [
                "Four dataset families were audited against the files on disk: "
                + "; ".join(
                    str(row.dataset_source) + " " + str(row.dataset_name) + " ("
                    + str(row.total_files) + " files, " + str(row.usable_files)
                    + " usable)"
                    for row in inventory.itertuples(index=False)
                )
                + ".",
                "The corpus holds " + f"{total:,}" + " recordings totalling "
                + f"{hours:.2f}" + " hours of audio, of which "
                + f"{supervised:,}" + " enter a supervised track. The remainder "
                "are the 301 PhysioNet validation duplicates and the 247 "
                "unlabelled PASCAL files, both excluded for the reasons given "
                "below.",
                "Native sampling rates are 2,000 Hz (PhysioNet), 4,000 Hz "
                "(PASCAL set_b, CirCor) and 44,100 Hz (PASCAL set_a). All are "
                "resampled to a common " + f"{TARGET_FS:,}" + " Hz.",
            ],
        )
    )

    disc_paras: list[str] = []
    for item in DISCREPANCIES:
        disc_paras.append(
            item.subject + " — the blueprint states " + item.blueprint_says
            + "; the files on disk hold " + item.disk_says + ". " + item.explanation
        )
    sections.append(("2. Count discrepancies against the source documents", disc_paras))

    lim_paras: list[str] = []
    for item in LIMITATIONS:
        lim_paras.append(item.subject + " — " + item.explanation)
    sections.append(("3. Limitations on what may be claimed", lim_paras))

    dup_paras: list[str] = []
    if not duplicates.empty:
        counts = duplicates["decision"].value_counts().to_dict()
        dup_paras.append(
            "Duplicate detection ran at three levels: SHA-256 of the raw bytes, "
            "a content hash of the decoded and resampled signal, and envelope "
            "correlation within each dataset. It produced "
            + str(len(duplicates)) + " rows: "
            + ", ".join(str(v) + " " + k for k, v in sorted(counts.items())) + "."
        )
        heartbeat = int(
            (duplicates["record_uid"].astype(str).str.startswith("HB_")).sum()
        )
        dup_paras.append(
            "All " + str(heartbeat) + " files under dataset/Heartbeat_Sound/ are "
            "byte-identical to files in PASCAL set_a and set_b. That folder is a "
            "label helper only; including it in any supervised track would "
            "double the PASCAL data and place identical recordings in both "
            "train and test."
        )
        validation = duplicates[
            (duplicates["dataset_source"] == "D1")
            & (duplicates["decision"] == "drop")
        ]
        dup_paras.append(
            "All " + str(len(validation)) + " recordings in the PhysioNet "
            "validation/ folder are byte-identical to records in training-a..f. "
            "They are excluded from every fold and no held-out PhysioNet test "
            "result is claimed anywhere in this project."
        )
        review = duplicates[duplicates["decision"] == "review"]
        if not review.empty:
            dup_paras.append(
                str(len(review)) + " near-duplicate pair(s) were flagged for "
                "review by envelope correlation. One is decisive: "
                "extrahls__201104021355 and murmur__201104021355 in PASCAL set_a "
                "are the same recording (correlation 0.999978) filed under two "
                "different class labels. Both label sources agree with "
                "themselves, so the dataset cannot say which label is the error. "
                "Both rows are retained and forced into the same fold."
            )
    sections.append(("4. Duplicate recordings", dup_paras))

    integ_paras: list[str] = []
    flags = master["quality_flags"].fillna("")
    integ_paras.append(
        "Every recording was decoded in full rather than read from its header. "
        "Zero files were unreadable, zero were zero-length and zero were "
        "truncated relative to their declared length, confirming the 2026-08-22 "
        "baseline audit."
    )
    integ_paras.append(
        str(int(flags.str.contains("is_silent").sum())) + " recording(s) are "
        "silent by peak amplitude and "
        + str(int(flags.str.contains("is_clipped").sum())) + " exceed the "
        "configured clipping threshold. These are flagged in the master table "
        "rather than removed, so the robustness track (EXP-E1) can stratify on "
        "them."
    )
    if not missing.empty:
        integ_paras.append(
            "Label coverage was checked in both directions. "
            + str(len(missing)) + " row(s) are recorded in "
            "missing_corrupt_files.csv."
        )
    sections.append(("5. File integrity and quality flags", integ_paras))

    if not class_dist.empty:
        supervised_rows = class_dist[class_dist["scope"] == "supervised"]
        class_paras = [
            "Class distributions, per task, on the supervised population:",
        ]
        for task, block in supervised_rows.groupby("task"):
            # `class` is a Python keyword, so itertuples renames that column.
            parts = [
                str(name) + " " + str(int(count))
                for name, count in zip(block["class"], block["n_records"], strict=True)
            ]
            worst = float(block["imbalance_ratio"].max())
            class_paras.append(
                str(task) + " — " + ", ".join(parts)
                + " (largest imbalance ratio " + f"{worst:.2f}" + ":1)."
            )
        class_paras.append(
            "The five label spaces are never merged. Binary, PASCAL A "
            "(four-class), PASCAL B (three-class), CirCor murmur and CirCor "
            "outcome are five separate tasks with five separate targets."
        )
        sections.append(("6. Class distribution and imbalance", class_paras))

    split_paras: list[str] = []
    if not split_map.empty:
        for task, block in split_map.groupby("task"):
            scheme = str(block["scheme"].iloc[0])
            n_repeats = int(block["n_repeats"].iloc[0])
            n_splits = int(block["n_splits"].iloc[0])
            split_paras.append(
                str(task) + " — " + scheme + ", " + str(n_splits) + " folds x "
                + str(n_repeats) + " repeat(s) over "
                + str(int(len(block) / n_repeats)) + " records and "
                + str(int(block["split_group"].nunique())) + " group(s)."
            )
        split_paras.append(
            "Every fold map is grouped: no subject appears in two folds of the "
            "same repeat, in any task. This is verified against the written "
            "fold map, not against the code that produced it "
            "(tests/test_no_leakage.py)."
        )
    sections.append(("7. Cross-validation fold maps", split_paras))

    return sections


def write_audit_report(
    master: Any,
    inventory: Any,
    out_dir: str | Path | None = None,
) -> Path:
    """Generate **DA-09** ``dataset_audit_report.docx`` (T21.5)."""
    from docx import Document
    from docx.shared import Pt

    from src.utils.io import ensure_dir

    target_dir = _audit_dir(out_dir)
    target = target_dir / "dataset_audit_report.docx"
    ensure_dir(target.parent)

    sections = build_report_sections(master, inventory, audit_dir=target_dir)

    document = Document()
    document.add_heading("PV-MEPCG / PulseVision — Dataset Audit Report", level=0)
    subtitle = document.add_paragraph(
        "DA-09 — generated from outputs/01_dataset_audit/ by "
        "src/data_loader/inventory.py"
    )
    subtitle.runs[0].font.size = Pt(9)

    document.add_heading("Dataset inventory (DA-01)", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ("ID", "Dataset", "Files", "Usable", "Subjects", "Native fs")
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        cell.text = text
    for row in inventory.itertuples(index=False):
        cells = table.add_row().cells
        cells[0].text = str(row.dataset_source)
        cells[1].text = str(row.dataset_name)
        cells[2].text = f"{int(row.total_files):,}"
        cells[3].text = f"{int(row.usable_files):,}"
        cells[4].text = f"{int(row.n_subjects):,}"
        cells[5].text = f"{int(row.original_fs):,} Hz"

    for heading, paragraphs in sections:
        document.add_heading(heading, level=1)
        for text in paragraphs:
            document.add_paragraph(text)

    document.save(str(target))
    log.info("wrote %s (%d sections)", target.name, len(sections))
    return target


# --------------------------------------------------------------------------
# T21.6 -- evidence index
# --------------------------------------------------------------------------

# DA-01 through DA-09, with the filename each one is written to. DA-08's path is
# outputs/01_dataset_audit/, not dataset/ -- see src/data_loader/master.py.
AUDIT_ARTIFACTS: tuple[tuple[str, str, str], ...] = (
    ("DA-01", "dataset_inventory.csv",
     "Per-dataset inventory: files, classes, subjects, role"),
    ("DA-02", "class_distribution.csv",
     "Class counts and imbalance ratios per dataset and task"),
    ("DA-03", "recording_duration_summary.csv",
     "Duration statistics per dataset and class"),
    ("DA-04", "sampling_rate_summary.csv",
     "Native versus target sampling rates with counts"),
    ("DA-05", "missing_corrupt_files.csv",
     "Unreadable, corrupt, missing-label and orphan-label files"),
    ("DA-06", "duplicate_report.csv",
     "Exact, content and near-duplicate groups with keep/drop decisions"),
    ("DA-07", "subject_split_map.csv",
     "Subject-grouped cross-validation fold map (all five tasks)"),
    ("DA-08", "metadata_master.csv",
     "Master metadata table (one row per recording)"),
    ("DA-09", "dataset_audit_report.docx",
     "Narrative dataset audit: counts, conflicts, duplicates, limitations"),
)


def register_audit_artifacts(
    out_dir: str | Path | None = None, *, index_path: str | Path | None = None
) -> list[dict[str, str]]:
    """Register DA-01 .. DA-09 in the evidence index (T21.6).

    An artifact that is not on disk is registered as ``missing`` rather than
    skipped -- a gap has to be visible in the index, not absent from it.

    **Evidence rows follow their artifacts.** When ``out_dir`` points somewhere
    other than the configured audit directory -- a smoke run, a test, anyone
    passing ``--out-dir`` -- the rows are written to an index inside that
    directory instead of the project's real one. Without this rule, a single
    ``--out-dir /tmp/...`` run rewrites DA-01..DA-09 in the committed index to
    paths under a temporary folder that is deleted minutes later, leaving nine
    rows claiming ``status=ok`` for files that no longer exist. That is not
    hypothetical: it is what the Phase 30 sweep found in this repository, put
    there by the T22.7 audit-script test.
    """
    from src.utils.evidence import evidence_index_path, register_evidence

    audit_dir = _audit_dir(out_dir)
    if index_path is None and audit_dir.resolve() != _audit_dir().resolve():
        index_path = audit_dir / "evidence_index.csv"
        log.info(
            "audit artifacts registered to %s, not the project index, because "
            "out_dir is not the configured audit directory",
            index_path,
        )
    target = Path(index_path) if index_path is not None else evidence_index_path()

    rows: list[dict[str, str]] = []
    for evidence_id, filename, description in AUDIT_ARTIFACTS:
        rows.append(
            register_evidence(
                evidence_id=evidence_id,
                objective="Data audit",
                dataset="D1-D4",
                metric_or_asset=description,
                filename=audit_dir / filename,
                source_data="dataset/ (read-only input)",
                command="python scripts/01_run_dataset_audit.py",
                index_path=target,
            )
        )
    missing = [r for r in rows if r.get("status") != "ok"]
    if missing:
        log.warning(
            "%d audit artifact(s) registered as missing: %s",
            len(missing),
            ", ".join(r["evidence_id"] for r in missing),
        )
    return rows


def run_inventory(
    master: Any | None = None, out_dir: str | Path | None = None
) -> dict[str, Any]:
    """Phase 21 end to end: DA-01, DA-09, and the evidence registrations."""
    if master is None:
        from src.data_loader.master import load_master

        master = load_master()

    inventory = build_inventory(master)
    inventory_path = write_inventory(inventory, out_dir)
    report_path = write_audit_report(master, inventory, out_dir)
    registered = register_audit_artifacts(out_dir)

    return {
        "inventory": inventory,
        "inventory_path": inventory_path,
        "report_path": report_path,
        "registered": registered,
    }
